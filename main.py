from fastapi import FastAPI, File, UploadFile, Request, BackgroundTasks
from fastapi.responses import FileResponse, HTMLResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
import os
import uuid
import time
import threading
import concurrent.futures
from gtts import gTTS
from pdfminer.high_level import extract_text
from docx import Document
import io
import math
import shutil
import hashlib
import sqlite3
import base64
from functools import lru_cache
from concurrent.futures import ThreadPoolExecutor
import multiprocessing
import traceback

# Número de procesadores disponibles para paralelización
NUM_CORES = max(1, multiprocessing.cpu_count() - 1)
MAX_WORKERS = NUM_CORES * 2  # Para operaciones I/O, podemos usar más workers que cores

app = FastAPI()

# Asegurarse de que los directorios necesarios existan
os.makedirs("audio", exist_ok=True)
os.makedirs("temp", exist_ok=True)
os.makedirs("templates", exist_ok=True)
os.makedirs("static", exist_ok=True)  # Asegurarse que la carpeta static existe
os.makedirs("cache", exist_ok=True)

# Montar carpeta estática para servir CSS, JS e imágenes
# IMPORTANTE: Estas rutas deben venir después de crear los directorios
app.mount("/static", StaticFiles(directory="static"), name="static")
app.mount("/audio", StaticFiles(directory="audio"), name="audio")

# Diccionario para almacenar el estado de las tareas
task_status = {}

# Inicializar la base de datos para el caché
def init_cache_db():
    conn = sqlite3.connect('cache/text_audio_cache.db')
    cursor = conn.cursor()
    cursor.execute('''
    CREATE TABLE IF NOT EXISTS text_chunks (
        hash_id TEXT PRIMARY KEY,
        text TEXT,
        audio_path TEXT,
        created_at INTEGER
    )
    ''')
    conn.commit()
    conn.close()

# Inicializar el caché
init_cache_db()

# Sistema de caché usando LRU para fragmentos de texto frecuentes
@lru_cache(maxsize=100)
def get_cached_audio_path(text_hash):
    """Busca un fragmento de texto en el caché y devuelve la ruta del audio si existe."""
    conn = sqlite3.connect('cache/text_audio_cache.db')
    cursor = conn.cursor()
    cursor.execute('SELECT audio_path FROM text_chunks WHERE hash_id = ?', (text_hash,))
    result = cursor.fetchone()
    conn.close()
    
    if result:
        return result[0]
    return None

def store_in_cache(text, audio_path):
    """Almacena un fragmento de texto y su audio correspondiente en el caché."""
    text_hash = hashlib.md5(text.encode('utf-8')).hexdigest()
    
    conn = sqlite3.connect('cache/text_audio_cache.db')
    cursor = conn.cursor()
    
    # Comprobar si ya existe
    cursor.execute('SELECT hash_id FROM text_chunks WHERE hash_id = ?', (text_hash,))
    if cursor.fetchone() is None:
        cursor.execute(
            'INSERT INTO text_chunks (hash_id, text, audio_path, created_at) VALUES (?, ?, ?, ?)',
            (text_hash, text, audio_path, int(time.time()))
        )
        conn.commit()
    
    conn.close()
    return text_hash

def clean_old_cache(max_age_days=30):
    """Limpia entradas de caché antiguas."""
    max_age = int(time.time()) - (max_age_days * 24 * 60 * 60)
    
    conn = sqlite3.connect('cache/text_audio_cache.db')
    cursor = conn.cursor()
    cursor.execute('SELECT audio_path FROM text_chunks WHERE created_at < ?', (max_age,))
    old_files = cursor.fetchall()
    
    # Eliminar archivos físicos
    for (file_path,) in old_files:
        if os.path.exists(file_path):
            try:
                os.remove(file_path)
            except:
                pass
    
    # Eliminar registros
    cursor.execute('DELETE FROM text_chunks WHERE created_at < ?', (max_age,))
    conn.commit()
    conn.close()

@app.get("/", response_class=HTMLResponse)
def read_root():
    """Servir la página HTML con el formulario."""
    try:
        # Verificar que el archivo HTML existe
        html_path = "templates/index.html"
        if not os.path.exists(html_path):
            return HTMLResponse(
                content=f"""
                <html>
                    <body>
                        <h1>Error: No se encontró el archivo index.html</h1>
                        <p>Por favor, asegúrate de que el archivo 'templates/index.html' existe.</p>
                    </body>
                </html>
                """,
                status_code=500
            )
            
        with open(html_path, "r", encoding="utf-8") as f:
            return HTMLResponse(content=f.read(), status_code=200)
    except Exception as e:
        print(f"Error al cargar la página: {str(e)}")
        traceback.print_exc()
        return HTMLResponse(
            content=f"""
            <html>
                <body>
                    <h1>Error al cargar la página: {str(e)}</h1>
                    <p>Detalles del error: {traceback.format_exc()}</p>
                </body>
            </html>
            """,
            status_code=500
        )

# Ruta de diagnóstico
@app.get("/health")
def health_check():
    """Comprueba el estado del servicio."""
    return {
        "status": "OK",
        "timestamp": time.time(),
        "dirs": {
            "templates": os.path.exists("templates"),
            "static": os.path.exists("static"),
            "audio": os.path.exists("audio"),
            "temp": os.path.exists("temp"),
            "cache": os.path.exists("cache")
        },
        "files": {
            "index.html": os.path.exists("templates/index.html"),
            "style.css": os.path.exists("static/style.css")
        }
    }


@app.post("/convert")
async def convert_file_to_audio(file: UploadFile = File(...), lang: str = "es"):
    """
    Recibe un archivo PDF o DOCX, extrae su texto y lo convierte en un archivo MP3.
    Procesa en segundo plano para archivos grandes.
    """
    try:
        # Identificar el tipo de archivo
        if not file.filename:
            return JSONResponse(content={"error": "No se ha proporcionado un archivo"}, status_code=400)
            
        file_ext = file.filename.split(".")[-1].lower()
        if file_ext not in ["pdf", "docx"]:
            return JSONResponse(content={"error": f"Formato de archivo no soportado: {file_ext}. Sube un PDF o DOCX."}, status_code=400)
        
        # Generar IDs únicos para los archivos
        task_id = str(uuid.uuid4().hex)
        temp_filename = f"temp/temp_{task_id}.{file_ext}"
        mp3_filename = f"audio/{task_id}.mp3"
        
        # Guardar el archivo subido
        file_content = await file.read()
        file_size = len(file_content)
        
        # Verificar tamaño máximo (50MB)
        if file_size > 50 * 1024 * 1024:
            return JSONResponse(content={"error": "El archivo excede el tamaño máximo permitido de 50MB"}, status_code=400)
        
        with open(temp_filename, "wb") as buffer:
            buffer.write(file_content)
        
        # Estimar tiempo basado en el tamaño del archivo (ahora más optimista)
        estimated_time = estimate_processing_time(file_size, file_ext)
        
        # Inicializar el estado de la tarea
        task_status[task_id] = {
            "status": "processing",
            "progress": 0,
            "estimated_time": estimated_time,
            "start_time": time.time(),
            "file_size": file_size
        }
        
        # Iniciar el procesamiento en un hilo separado    
        thread = threading.Thread(
        target=process_file_thread,
        args=(task_id, temp_filename, file_ext, mp3_filename, lang)
        )
        thread.daemon = True
        thread.start()
    
        return JSONResponse(content={
            "task_id": task_id,
            "estimated_time": estimated_time,
            "file_size": file_size
        })
        
    except Exception as e:
        print(f"Error al procesar la solicitud de conversión: {str(e)}")
        traceback.print_exc()
        return JSONResponse(content={"error": str(e)}, status_code=500)

@app.get("/task/{task_id}")
async def get_task_status(task_id: str):
    """Obtener el estado actual de una tarea de procesamiento."""
    if task_id not in task_status:
        return JSONResponse(content={"error": "Tarea no encontrada"}, status_code=404)
    
    status_info = task_status[task_id].copy()
    
    # Si la tarea está completa, devolver también las URLs
    if status_info["status"] == "completed":
        status_info["audio_url"] = f"/audio/{task_id}.mp3"
        
        # Limpiar el estado después de completado y reportado
        if "text" in status_info:
            text = status_info["text"]
            del status_info["text"]  # No enviar todo el texto en cada consulta de estado
        else:
            # Si el texto ya fue eliminado, leer el archivo de texto
            try:
                with open(f"temp/text_{task_id}.txt", "r", encoding="utf-8") as f:
                    text = f.read()
            except:
                text = "El texto extraído no está disponible"
                
        status_info["text"] = text
    
    # Actualizar el tiempo transcurrido
    if status_info["status"] == "processing":
        elapsed = time.time() - status_info["start_time"]
        status_info["elapsed_time"] = elapsed
        
        # Actualizar la estimación restante
        if status_info["progress"] > 0:
            remaining = (elapsed / status_info["progress"]) * (100 - status_info["progress"])
            status_info["remaining_time"] = remaining
    
    return JSONResponse(content=status_info)

def estimate_processing_time(file_size, file_ext):
    """Estima el tiempo de procesamiento basado en el tamaño del archivo."""
    # Tiempo base más bajo gracias a la optimización
    base_time = 5  # segundos base
    
    # Factores de tiempo reducidos por paralelización
    factors = {
        "pdf": 0.8,  # segundos por MB para PDF (reducido por procesamiento paralelo)
        "docx": 0.6  # segundos por MB para DOCX (reducido por procesamiento paralelo)
    }
    
    # Calcular tiempo estimado
    size_mb = file_size / (1024 * 1024)
    factor = factors.get(file_ext, 1.0)
    
    # La estimación ahora considera que el procesamiento es paralelo
    estimated_time = base_time + (size_mb * factor / NUM_CORES)
    
    # Añadir factor para archivos muy grandes (pero menor que antes)
    if size_mb > 20:
        estimated_time *= 1.1
    
    return math.ceil(estimated_time)

# Árbol de fragmentos para optimizar el procesamiento de texto
class TextChunkTree:
    """
    Estructura de árbol para organizar fragmentos de texto y optimizar la conversión.
    Esta estructura ayuda a identificar patrones comunes y frases repetitivas.
    """
    def __init__(self):
        self.chunks = {}
        self.children = {}
        
    def add_chunk(self, text, chunk_id):
        """Añade un fragmento al árbol"""
        # Usar las primeras palabras como clave
        words = text.split()
        key = " ".join(words[:3]) if len(words) >= 3 else text[:20]
        
        if key not in self.children:
            self.children[key] = TextChunkTree()
        
        self.children[key].chunks[chunk_id] = text
    
    def find_similar_chunks(self, text, threshold=0.8):
        """Encuentra fragmentos similares en el árbol"""
        words = text.split()
        key = " ".join(words[:3]) if len(words) >= 3 else text[:20]
        
        similar_chunks = []
        
        if key in self.children:
            # Calcular similitud con fragmentos existentes
            for chunk_id, chunk_text in self.children[key].chunks.items():
                similarity = self._calculate_similarity(text, chunk_text)
                if similarity >= threshold:
                    similar_chunks.append((chunk_id, similarity))
        
        return similar_chunks
    
    def _calculate_similarity(self, text1, text2):
        """Calcula la similitud entre dos textos (método simple)"""
        # Convertir a conjuntos de palabras
        words1 = set(text1.lower().split())
        words2 = set(text2.lower().split())
        
        # Usar coeficiente de Jaccard
        if not words1 or not words2:
            return 0
        
        intersection = words1.intersection(words2)
        union = words1.union(words2)
        
        return len(intersection) / len(union)

# Crear un árbol de fragmentos global
text_chunk_tree = TextChunkTree()

def extract_text_from_pdf_optimized(pdf_path: str) -> str:
    """Extrae el texto del PDF usando pdfminer.six con mayor robustez y diagnóstico."""
    try:
        print(f"Intentando extraer texto de PDF: {pdf_path}")
        
        # Verificar si el archivo existe y tiene contenido
        if not os.path.exists(pdf_path):
            print(f"Error: El archivo PDF no existe: {pdf_path}")
            return ""
            
        if os.path.getsize(pdf_path) == 0:
            print(f"Error: El archivo PDF está vacío: {pdf_path}")
            return ""
        
        # Intentar con parámetros más básicos primero
        try:
            text = extract_text(pdf_path, page_numbers=None, maxpages=0)
            if text.strip():
                return text
        except Exception as e:
            print(f"Primer intento de extracción de PDF falló: {str(e)}")
        
        # Si falló, intentar con configuración alternativa
        try:
            text = extract_text(
                pdf_path, 
                page_numbers=None, 
                maxpages=0, 
                laparams=None  # Sin parámetros de layout
            )
            return text
        except Exception as e:
            print(f"Segundo intento de extracción de PDF falló: {str(e)}")
            
        # Último intento con configuración mínima
        from pdfminer.high_level import extract_text_to_fp
        from pdfminer.layout import LAParams
        from io import StringIO
        
        output = StringIO()
        with open(pdf_path, 'rb') as fp:
            extract_text_to_fp(fp, output, laparams=LAParams(), 
                             output_type='text', codec='utf-8')
        return output.getvalue()
            
    except Exception as e:
        print(f"Error al extraer texto de PDF (método final): {str(e)}")
        traceback.print_exc()  # Imprimir traza completa para diagnóstico
        return "No se pudo extraer texto del PDF. Por favor, verifique que el archivo no esté protegido o dañado."

def extract_text_from_docx_optimized(docx_path: str) -> str:
    """Extrae el texto de un archivo DOCX con mayor robustez y diagnóstico."""
    try:
        print(f"Intentando extraer texto de DOCX: {docx_path}")
        
        # Verificar si el archivo existe y tiene contenido
        if not os.path.exists(docx_path):
            print(f"Error: El archivo DOCX no existe: {docx_path}")
            return ""
            
        if os.path.getsize(docx_path) == 0:
            print(f"Error: El archivo DOCX está vacío: {docx_path}")
            return ""
        
        # Intentar abrir el documento
        try:
            doc = Document(docx_path)
        except Exception as e:
            print(f"Error al abrir DOCX con python-docx: {str(e)}")
            # Intentar método alternativo
            return extract_docx_fallback(docx_path)
        
        text_parts = []
        
        # Extraer texto de párrafos
        try:
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
        except Exception as e:
            print(f"Error al extraer párrafos: {str(e)}")
        
        # Extraer texto de tablas
        try:
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        text_parts.append(" | ".join(row_text))
        except Exception as e:
            print(f"Error al extraer tablas: {str(e)}")
        
        result = "\n".join(text_parts)
        
        # Si no se extrajo nada, probar método alternativo
        if not result.strip():
            print("No se extrajo texto con python-docx, intentando método alternativo")
            return extract_docx_fallback(docx_path)
            
        return result
        
    except Exception as e:
        print(f"Error al extraer texto de DOCX (general): {str(e)}")
        traceback.print_exc()  # Imprimir traza completa para diagnóstico
        return "No se pudo extraer texto del DOCX. Por favor, verifique que el archivo no esté protegido o dañado."

def extract_docx_fallback(docx_path):
    """Método alternativo para extraer texto de DOCX usando zipfile."""
    try:
        print("Usando método alternativo para DOCX con zipfile")
        import zipfile
        from xml.etree.ElementTree import XML
        
        # Definir NAMESPACE
        NAMESPACE = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        PARA = NAMESPACE + 'p'
        TEXT = NAMESPACE + 't'
        
        # Abrir el archivo como un zip
        with zipfile.ZipFile(docx_path) as docx:
            # Extraer document.xml que contiene el contenido
            try:
                content = docx.read('word/document.xml')
            except KeyError:
                print("No se pudo encontrar 'word/document.xml' en el archivo DOCX")
                return ""
                
            # Analizar el XML
            tree = XML(content)
            
            # Extraer párrafos
            paragraphs = []
            for paragraph in tree.iter(PARA):
                texts = [node.text for node in paragraph.iter(TEXT) if node.text]
                if texts:
                    paragraphs.append(''.join(texts))
                    
            return '\n'.join(paragraphs)
    except Exception as e:
        print(f"Error en método alternativo DOCX: {str(e)}")
        return ""

def split_text_optimized(text: str, max_length: int = 800) -> list:
    """
    Divide el texto en fragmentos más pequeños para procesamiento eficiente.
    Utiliza un tamaño máximo más pequeño para mejor paralelización.
    """
    # Si el texto es más corto que el máximo, devolverlo como un solo fragmento
    if len(text) <= max_length:
        return [text]
    
    # Dividir por párrafos
    paragraphs = text.split('\n')
    
    # Agrupar párrafos en chunks de tamaño óptimo
    chunks = []
    current_chunk = ""
    
    for para in paragraphs:
        if not para.strip():
            continue
            
        if len(current_chunk) + len(para) + 1 > max_length and current_chunk:
            chunks.append(current_chunk)
            current_chunk = para
        else:
            if current_chunk:
                current_chunk += '\n' + para
            else:
                current_chunk = para
    
    # Añadir el último chunk
    if current_chunk:
        chunks.append(current_chunk)
    
    # Si hay párrafos muy largos, dividirlos por oraciones
    final_chunks = []
    for chunk in chunks:
        if len(chunk) <= max_length:
            final_chunks.append(chunk)
        else:
            # Dividir por oraciones
            sentences = split_into_sentences(chunk)
            current_sentence_chunk = ""
            
            for sentence in sentences:
                if len(current_sentence_chunk) + len(sentence) > max_length and current_sentence_chunk:
                    final_chunks.append(current_sentence_chunk)
                    current_sentence_chunk = sentence
                else:
                    if current_sentence_chunk:
                        current_sentence_chunk += ' ' + sentence
                    else:
                        current_sentence_chunk = sentence
            
            if current_sentence_chunk:
                final_chunks.append(current_sentence_chunk)
    
    return final_chunks

def split_into_sentences(text):
    """Divide un texto en oraciones."""
    # Patrones comunes de final de oración en español
    sentence_endings = ['. ', '! ', '? ', '.\n', '!\n', '?\n', '... ', '...\n']
    result = []
    
    remaining_text = text
    while remaining_text:
        min_idx = len(remaining_text)
        min_ending = None
        
        for ending in sentence_endings:
            idx = remaining_text.find(ending)
            if idx != -1 and idx < min_idx:
                min_idx = idx
                min_ending = ending
        
        if min_ending:
            sentence = remaining_text[:min_idx + len(min_ending)]
            result.append(sentence)
            remaining_text = remaining_text[min_idx + len(min_ending):]
        else:
            result.append(remaining_text)
            break
    
    return result

def process_file_thread(task_id, file_path, file_ext, mp3_filename, lang: str = "es"):
    try:
        if not os.path.exists(file_path):
            task_status[task_id]["status"] = "error"
            task_status[task_id]["error"] = f"El archivo no existe: {file_path}"
            return
            
        if os.path.getsize(file_path) == 0:
            task_status[task_id]["status"] = "error"
            task_status[task_id]["error"] = "El archivo está vacío"
            return
            
        task_status[task_id]["progress"] = 5
        print(f"Iniciando extracción de texto del archivo {file_ext}: {file_path}")
        
        try:
            if file_ext == "pdf":
                text = extract_text_from_pdf_optimized(file_path)
            else:
                text = extract_text_from_docx_optimized(file_path)
                
            print(f"Texto extraído: {len(text)} caracteres")
            task_status[task_id]["progress"] = 30
            
            if not text or not text.strip():
                task_status[task_id]["status"] = "error"
                task_status[task_id]["error"] = "No se pudo extraer texto del archivo. Verifique que no esté protegido o dañado."
                cleanup_temp_files([file_path])
                return
            
            text_filename = f"temp/text_{task_id}.txt"
            with open(text_filename, "w", encoding="utf-8") as f:
                f.write(text)
            task_status[task_id]["text"] = text
            
        except Exception as e:
            task_status[task_id]["status"] = "error"
            task_status[task_id]["error"] = f"Error al extraer texto: {str(e)}"
            traceback.print_exc()
            cleanup_temp_files([file_path])
            return
        
        try:
            text_chunks = split_text_optimized(text)
            print(f"Texto dividido en {len(text_chunks)} fragmentos")
            task_status[task_id]["progress"] = 40
            
            # Aquí usamos la versión paralela y pasamos el idioma seleccionado
            audio_files = process_chunks_parallel(text_chunks, task_id, lang=lang)
            
            task_status[task_id]["progress"] = 90
            
            if audio_files:
                try:
                    concatenate_audio_files_simple(audio_files, mp3_filename)
                    print(f"Archivos de audio concatenados: {len(audio_files)}")
                except Exception as e:
                    print(f"Error al concatenar audio: {str(e)}")
                    fallback_concatenate(audio_files, mp3_filename)
            else:
                print("No se generaron archivos de audio para concatenar")
                with open(mp3_filename, 'wb') as f:
                    f.write(b'')
            
            cleanup_temp_files(audio_files + [file_path])
            
            task_status[task_id]["progress"] = 100
            task_status[task_id]["status"] = "completed"
            task_status[task_id]["completion_time"] = time.time()
                
        except Exception as e:
            print(f"Error en el procesamiento de archivo: {str(e)}")
            traceback.print_exc()
            task_status[task_id]["status"] = "error"
            task_status[task_id]["error"] = f"Error en procesamiento: {str(e)}"
            cleanup_temp_files([file_path])
    
    except Exception as e:
        print(f"Error general en procesamiento: {str(e)}")
        traceback.print_exc()
        task_status[task_id]["status"] = "error"
        task_status[task_id]["error"] = f"Error: {str(e)}"
        if os.path.exists(file_path):
            try:
                os.remove(file_path)
            except:
                pass


def text_to_speech_optimized(text: str, output_filename: str, lang: str = "es"):
    """Convierte texto a voz usando gTTS con manejo de errores y soporte para idioma."""
    try:
        if not text or not text.strip():
            print("Error: Texto vacío para conversión a voz")
            with open(output_filename, 'wb') as f:
                f.write(b'')
            return
            
        text = text.replace('|', ',').replace('\x00', ' ')
        
        if len(text) > 5000:
            print(f"Texto demasiado largo ({len(text)} caracteres), recortando a 5000")
            text = text[:5000]
        
        try:
            print(f"Intentando convertir texto a voz ({len(text)} caracteres) en idioma '{lang}'")
            tts = gTTS(text=text, lang=lang, slow=False)
            tts.save(output_filename)
            print(f"Audio guardado: {output_filename}")
        except AssertionError as e:
            print(f"Error de aserción en gTTS: {str(e)}")
            with open(output_filename, 'wb') as f:
                f.write(b'')
        except Exception as e:
            print(f"Error desconocido en gTTS: {str(e)}")
            if len(text) > 100:
                print("Intentando con fragmento del texto")
                try:
                    tts = gTTS(text=text[:100] + "...", lang=lang, slow=False)
                    tts.save(output_filename)
                except:
                    with open(output_filename, 'wb') as f:
                        f.write(b'')
            else:
                with open(output_filename, 'wb') as f:
                    f.write(b'')
                    
    except Exception as e:
        print(f"Error general en text-to-speech: {str(e)}")
        traceback.print_exc()
        with open(output_filename, 'wb') as f:
            f.write(b'')


def concatenate_audio_files_simple(input_files, output_file):
    """Versión simplificada de concatenación de archivos para diagnóstico."""
    with open(output_file, 'wb') as outfile:
        for fname in input_files:
            if os.path.exists(fname) and os.path.getsize(fname) > 0:
                with open(fname, 'rb') as infile:
                    outfile.write(infile.read())
                    
def fallback_concatenate(input_files, output_file):
    """Método de último recurso para concatenar archivos."""
    # Simplemente copiar el primer archivo si existe
    for file in input_files:
        if os.path.exists(file) and os.path.getsize(file) > 0:
            shutil.copy2(file, output_file)
            print(f"Método de reserva: copiando {file} a {output_file}")
            return
            
    # Si no hay archivos, crear uno vacío
    with open(output_file, 'wb') as f:
        f.write(b'')

def cleanup_temp_files(file_paths):
    """Elimina archivos temporales de forma segura."""
    for file_path in file_paths:
        if os.path.exists(file_path):
            try:
                os.remove(file_path)
            except:
                pass

def process_chunks_parallel(text_chunks, task_id, lang: str = "es"):
    """Procesa múltiples fragmentos de texto en paralelo usando un ThreadPool y soporte de idioma."""
    audio_files = []
    
    def process_chunk(chunk_data):
        idx, chunk = chunk_data
        chunk_filename = f"temp/chunk_{task_id}_{idx}.mp3"
        
        # Calcular hash y buscar en caché (sin cambios)
        chunk_hash = hashlib.md5(chunk.encode('utf-8')).hexdigest()
        cached_path = get_cached_audio_path(chunk_hash)
        if cached_path and os.path.exists(cached_path):
            shutil.copy2(cached_path, chunk_filename)
        else:
            similar_chunks = text_chunk_tree.find_similar_chunks(chunk)
            if similar_chunks:
                most_similar = max(similar_chunks, key=lambda x: x[1])
                similar_id, similarity = most_similar
                similar_path = f"temp/chunk_{similar_id}.mp3"
                if os.path.exists(similar_path) and similarity > 0.9:
                    shutil.copy2(similar_path, chunk_filename)
                else:
                    text_to_speech_optimized(chunk, chunk_filename, lang=lang)
                    store_in_cache(chunk, chunk_filename)
            else:
                text_to_speech_optimized(chunk, chunk_filename, lang=lang)
                store_in_cache(chunk, chunk_filename)
                text_chunk_tree.add_chunk(chunk, f"{task_id}_{idx}")
        
        return chunk_filename
    
    with ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
        chunk_data = [(i, chunk) for i, chunk in enumerate(text_chunks)]
        batch_size = 10
        total_chunks = len(chunk_data)
        audio_files = []
        
        for i in range(0, total_chunks, batch_size):
            batch = chunk_data[i:min(i+batch_size, total_chunks)]
            batch_results = list(executor.map(process_chunk, batch))
            audio_files.extend(batch_results)
            
            progress = 40 + (50 * min(i + batch_size, total_chunks) / total_chunks)
            task_status[task_id]["progress"] = progress
    
    return audio_files


# Limpieza periódica mejorada
def cleanup_thread():
    """Función para limpiar periódicamente las tareas y archivos antiguos."""
    while True:
        try:
            current_time = time.time()
            to_delete = []
            
            # Hacer una copia del diccionario para evitar errores
            tasks_copy = task_status.copy()
            
            for task_id, task_info in tasks_copy.items():
                # Limpiar tareas completadas después de 1 hora
                if task_info["status"] in ["completed", "error"]:
                    if "completion_time" in task_info and (current_time - task_info["completion_time"]) > 3600:
                        to_delete.append(task_id)
            
            # Eliminar tareas antiguas
            for task_id in to_delete:
                if task_id in task_status:
                    del task_status[task_id]
                
                # Eliminar archivos de texto asociados
                text_file = f"temp/text_{task_id}.txt"
                if os.path.exists(text_file):
                    try:
                        os.remove(text_file)
                    except:
                        pass
            
            # Limpieza de caché periódica (cada 5 ejecuciones)
            if time.time() % 5 == 0:
                clean_old_cache()
            
            # Dormir durante 5 minutos
            time.sleep(300)
            
        except Exception as e:
            print(f"Error en el hilo de limpieza: {str(e)}")
            time.sleep(60)

@app.on_event("startup")
def startup_event():
    """Iniciar el hilo de limpieza al arrancar la aplicación."""
    cleanup_thread_instance = threading.Thread(target=cleanup_thread)
    cleanup_thread_instance.daemon = True
    cleanup_thread_instance.start()